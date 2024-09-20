from docx import Document


def read_docx(file_path):
    document = Document(file_path)
    data = []

    for table in document.tables:

        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if len(cell_text) > 0:
                    row_data.append(cell_text)

            data.append(row_data)

    return data


def process_data(df):
    processed_data = []
    previous_sandhi = ""
    previous_padya = ""
    previous_tipanni = ""

    for index, row in df.iterrows():
        sandhi, padya, tipanni = row['ಸಂಧಿ'], row['ಪದ್ಯ'], row['ಟಿಪ್ಪಣಿ']
        if padya:
            if previous_tipanni:
                processed_data.append([previous_sandhi, previous_padya, previous_tipanni])
            previous_sandhi = sandhi
            previous_padya = padya
            previous_tipanni = tipanni
        else:
            previous_tipanni += "\n" + tipanni

    if previous_tipanni:
        processed_data.append([previous_sandhi, previous_padya, previous_tipanni])

    return pd.DataFrame(processed_data, columns=['ಸಂಧಿ', 'ಪದ್ಯ', 'ಟಿಪ್ಪಣಿ'])


def extract_tippani():
    file_path = '../docs/ಟಿಪ್ಪಣಿ-ಸಭಾಪರ್ವ-ಅರಣ್ಯಪರ್ವ-ವಿರಾಟಪರ್ವ.docx'
    data = read_docx(file_path)

    # Filter rows to only include those with exactly 3 columns
    filtered_data = [entry for entry in data]
    print(filtered_data)
    # Create a DataFrame
    df = pd.DataFrame(filtered_data, columns=['ಸಂಧಿ', 'ಪದ್ಯ', 'ಗುಣಗಳು', 'ಟಿಪ್ಪಣಿ', ''])

    # Forward fill the 'ಸಂಧಿ' column for any empty cells
    df['ಸಂಧಿ'] = df['ಸಂಧಿ'].replace('', pd.NA).ffill()

    # Process the DataFrame
    processed_df = process_data(df)

    # Save to CSV
    csv_file_path = 'all_tippani.csv'
    processed_df.to_csv(csv_file_path, index=False, encoding='utf-8-sig')

    print(f'Data successfully saved to {csv_file_path}')


import docx
import pandas as pd


def extract_parva_data_from_table(docx_file):
    doc = docx.Document(docx_file)

    data = []
    current_parva = None
    last_sandhi = None
    last_padya = None
    last_tippani = None

    # Iterate over tables in the document
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            # Extract text from each cell in the row
            cells = [cell.text.strip() for cell in row.cells]

            # Check if the first row is the Parva name (only one cell with text)
            if str(cells[0]).endswith('ಪರ್ವ'):
                current_parva = cells[0]
                print('current parva: ', current_parva)
                continue

            # Skip the header row if it has the terms 'ಸಂಧಿ', 'ಪದ್ಯ', and 'ಟಿಪ್ಪಣಿ'
            if "ಸಂಧಿ" in cells and "ಪದ್ಯ" in cells and "ಟಿಪ್ಪಣಿ" in cells:
                continue

            # Check if the row contains actual data (sandhi, padya, tippani)
            if len(cells) >= 3:
                sandhi = cells[0].strip() if cells[0] else last_sandhi
                padya = cells[1].strip() if cells[1] else last_padya
                tippani = cells[2].strip() if len(cells) > 2 else ""
                # Append the row to the data
                if not str(sandhi).startswith('ಸಂಧಿ'):
                    sandhi = 'ಸಂಧಿ ' + str(sandhi).strip()
                padya = str(padya).replace('ಪದ್ಯ', '').strip()
                current_parva = str(current_parva).strip()
                data.append([current_parva, sandhi, padya, tippani])

                # Update the last values
                last_sandhi = sandhi
                last_padya = padya
                last_tippani = tippani

    return data


def write_to_csv(data, output_file):
    # Convert to DataFrame and write to CSV
    df = pd.DataFrame(data, columns=['Parva', 'Sandhi', 'Padya', 'Tippani'])
    df.to_csv(output_file, index=False, encoding='utf-8')


def main(docx_file, output_csv):
    data = extract_parva_data_from_table(docx_file)
    write_to_csv(data, output_csv)
    print(f"Data has been successfully written to {output_csv}")


# Specify your input .docx file and the desired output CSV file
file_path = '../docs/ಟಿಪ್ಪಣಿ-ಸಭಾಪರ್ವ-ಅರಣ್ಯಪರ್ವ-ವಿರಾಟಪರ್ವ.docx'
output_csv = 'output.csv'

main(file_path, output_csv)
