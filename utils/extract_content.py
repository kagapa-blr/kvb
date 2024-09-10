import pandas as pd
from docx import Document

KANNADA_DIGITS = {'೦', '೧', '೨', '೩', '೪', '೫', '೬', '೭', '೮', '೯'}


def read_docx_for_paragraphs(file_path):
    doc = Document(file_path)
    paragraphs = []
    current_paragraph = ""

    for para in doc.paragraphs:
        text = para.text.strip()

        if text:  # If the line is not empty or has a length greater than 0
            if current_paragraph:
                current_paragraph += "\n" + text  # Append the text to the current paragraph
            else:
                current_paragraph = text
        else:
            if current_paragraph:
                paragraphs.append(current_paragraph)  # Append the accumulated paragraph to the list
                current_paragraph = ""  # Reset for the next paragraph

    # Add the last paragraph if not empty
    if current_paragraph:
        paragraphs.append(current_paragraph)

    return paragraphs


def starts_with_kannada_digit(text):
    return text and text[0] in KANNADA_DIGITS


def read_docx(file_path):
    doc = Document(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        paragraphs.append(para.text)
    return paragraphs


def classify_paragraph(paragraph):
    """Classify the paragraph based on the given rules."""
    text = paragraph.strip()
    if text == "ಸಭಾಪರ್ವ":
        return "parva"
    elif text.startswith("ಸಂಧಿ") or text.endswith("ಸಂಧಿ"):
        return "sandhi"
    elif text.endswith("||"):
        return "padya"
    elif starts_with_kannada_digit(text):
        return "gadya"
    elif text.startswith("ಪಾಠಾಂತರ"):
        return "patantar"
    elif text.startswith("ಟಿಪ್ಪಣಿ"):
        return "tippani"
    elif text.startswith("ಸೂಚನೆ") or text.startswith("ಸೂ"):
        return "suchane"
    elif text.startswith("ಅರ್ಥ"):
        return "artha"
    else:
        return "unknown"


def read_and_classify_paragraphs(docx_file):
    all_paragraphs = read_docx_for_paragraphs(docx_file)  # read_docx(docx_file)
    classified_paragraphs = []
    current_entry = {"parva": "", "sandhi": "", "padya": "", "gadya": "", "patantar": "", "tippani": "", "suchane": "",
                     "artha": ""}
    for para in all_paragraphs:
        para = para.strip()
        # print('para: ', para, type(para), len(para))
        section = classify_paragraph(para)

        # Update the current entry with the identified section
        current_entry[section] = para

        # When the section is "artha", add the current entry to the list and reset it
        if section == "artha":
            classified_paragraphs.append(current_entry)
            current_entry = {"parva": "", "sandhi": "", "padya": "", "gadya": "", "patantar": "", "tippani": "",
                             "suchane": "", "artha": ""}

    # Add the last accumulated entry if not empty
    if any(current_entry.values()):
        classified_paragraphs.append(current_entry)

    return classified_paragraphs


def parva_extract():
    # Specify the path to your DOCX file
    docx_file_path = '../docs/ಸಭಾಪರ್ವ-new.docx'

    # Call the function to read and classify paragraphs
    classified_paragraphs = read_and_classify_paragraphs(docx_file_path)

    # Convert the list of dictionaries to a DataFrame
    df = pd.DataFrame(classified_paragraphs)

    # Fill 'parva' column with the constant value 'parva' for all rows
    df['parva'] = 'ಸಭಾಪರ್ವ'

    # Propagate the 'sandhi' value to subsequent rows if the current 'sandhi' is empty
    df['sandhi'] = df['sandhi'].replace('', pd.NA).ffill()
    df = df[df['padya'].notna() & (df['padya'].str.strip() != '')]
    # df = df[(df['padya'].str.strip() == '')]

    # Save the DataFrame to a CSV file
    csv_file_path = 'ಸಭಾಪರ್ವ-new.csv'
    df.to_csv(csv_file_path, index=False, encoding='utf-8')

    print(f"Classified paragraphs have been saved to {csv_file_path}")
